from docx.document import Document
import datetime
from app.core.models import ProtocolData
from app.utils.logger import logger

class BaseSection:
    '''Базовый класс для всех разделов протокола'''
    
    def __init__(self, document: Document, protocol_data: ProtocolData):
        """
        Инициализация базового раздела
        :param document: Объект документа Word  
        :param protocol_data: Данные протокола
        """
        self.document = document
        self.protocol_data = protocol_data
        self.logger = logger.bind(section=self.__class__.__name__)
        
        self.logger.debug("Инициализация раздела протокола")
        
    def build(self):
        """Основной метод для построения раздела (должен быть переопределен)"""
        raise NotImplementedError("Метод build() должен быть реализован в подклассе")
    
    def _add_heading(self, text: str, level: int = 1):
        """Добавляет заголовок в документ"""
        self.logger.info(f"Добавление заголовка: {text}")
        self.document.add_heading(text, level=level)
        
    def _add_paragraph(self, text: str, bold = False):
        """Добавляет абзац текста в документ"""
        self.logger.debug(f"Добавление абзаца: {text[:50]}...")
        paragraph = self.document.add_paragraph(text)
        if bold:
            paragraph.runs[0].bold = True
            
    def _add_table(self, rows: int, cols: int):
        """Добавляет таблицу в документ"""
        self.logger.debug(f"Создание таблицы: {rows}x{cols}")
        return self.document.add_table(rows=rows, cols=cols)
    
    def _formated_date(self, date: datetime) -> str:
        """Форматирование даты в российский формат"""
        return date.strftime("%d.%m.%Y")
        