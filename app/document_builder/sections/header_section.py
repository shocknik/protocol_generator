from docx.shared import Pt, Cm
from .base_section import BaseSection
from config import TEST_CENTER, TEST_CENTER_ADDRESS, CERTIFICATION_TEXT, TestType

class HeaderSection(BaseSection):
    '''Титульник'''
    
    def build(self):
        self.logger.info("Начало построения титульного листа")
        
        self._add_heading("ПРОТОКОЛ № __", level=0)
        
        center_info = (
            f'Испытательный центр: {TEST_CENTER}'
            f'Адрес: {TEST_CENTER_ADDRESS}'
        )
        self._add_paragraph(center_info)
        
        # Добавляем таблицу для титульника
        table = self._add_table(rows=3, cols=3)
        
        # Настройка размеров колонок
        table.autofit = False
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(15)
        table.columns[2].width = Cm(5)
        
        # Заполняем ячейки
        self._fill_header_table(table)
        
        self.logger.success("Титульный лист сформирован")
        
    def _fill_header_table(self, table):
        """Метод заполнения таблицы титульника"""
        
        row0 = table.rows[0]
        row0.cells[0].text = f"Протокол № {self.protocol_data.test_object.id}"
        row0.cells[1].text = f"от {self._formated_date(self.protocol_data.test_dates.end_date)}"
        
        
        # Утверждающая подпись
        row0.cells[0]
        approval_cell = row0.cells[2]
        approval_cell.text = "УТВЕРЖДАЮ"
        approval_cell.add_paragraph(CERTIFICATION_TEXT)
        approval_cell.add_paragraph("____________________")
        approval_cell.add_paragraph(self._formated_date(self.protocol_data.test_dates.end_date))

        

        